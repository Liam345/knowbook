"""
DOCX processor for KnowBook.

Handles Word documents (.docx) using python-docx library.
Preserves document structure including headings, lists, and tables.
"""
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List

from app.utils.text import (
    build_and_save_processed_output,
    count_tokens
)
from app.services.embedding_service import embedding_service
from app.services.summary_service import summary_service


def process_docx(
    project_id: str,
    source_id: str,
    source: Dict[str, Any],
    raw_file_path: Path,
    source_service
) -> Dict[str, Any]:
    """
    Process a Word document.

    Pipeline:
    1. Extract text from DOCX using python-docx
    2. Preserve structure (headings, lists, tables)
    3. Create processed output with page markers
    4. Generate embeddings
    5. Generate summary

    Args:
        project_id: Project UUID
        source_id: Source UUID
        source: Source metadata dictionary
        raw_file_path: Path to the raw DOCX file
        source_service: Source service for status updates

    Returns:
        Result dictionary with success status and info
    """
    try:
        from docx import Document

        # Load document
        doc = Document(str(raw_file_path))

        # Extract content with structure preservation
        content_parts = []
        paragraph_count = 0
        table_count = 0

        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                paragraph_count += 1
                text = _extract_paragraph_text(element, doc)
                if text.strip():
                    content_parts.append(text)

            # Handle tables
            elif element.tag.endswith('tbl'):
                table_count += 1
                table_text = _extract_table_text(element, doc)
                if table_text.strip():
                    content_parts.append(table_text)

        content = '\n\n'.join(content_parts)

        if not content.strip():
            return {
                "success": False,
                "error": "Document is empty or contains no extractable text"
            }

        # Count tokens and characters
        token_count = count_tokens(content)
        character_count = len(content)

        # Build and save processed output
        source_name = source.get('name', source.get('original_name', 'unknown'))
        metadata = {
            'character_count': character_count,
            'token_count': token_count,
            'paragraph_count': paragraph_count,
            'table_count': table_count
        }

        data_dir = Path(__file__).parent.parent.parent / 'data'
        build_and_save_processed_output(
            project_id=project_id,
            source_id=source_id,
            pages=[content],  # Single page for DOCX (no natural page breaks)
            source_type='DOCX',
            source_name=source_name,
            metadata=metadata,
            data_dir=data_dir
        )

        processing_info = {
            'character_count': character_count,
            'token_count': token_count,
            'paragraph_count': paragraph_count,
            'table_count': table_count,
            'pages_processed': 1,
            'processed_at': datetime.utcnow().isoformat()
        }

        # Update status to embedding
        source_service._update_source_status(
            project_id, source_id, 'embedding',
            processing_info=processing_info
        )

        # Process embeddings
        embedding_info = embedding_service.process_embeddings(
            project_id=project_id,
            source_id=source_id,
            source_name=source_name
        )

        # Generate summary
        summary_info = summary_service.generate_summary(
            project_id=project_id,
            source_id=source_id,
            source_name=source_name
        )

        return {
            "success": True,
            "processing_info": processing_info,
            "embedding_info": embedding_info,
            "summary_info": summary_info
        }

    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


def _extract_paragraph_text(element, doc) -> str:
    """
    Extract text from a paragraph element with formatting hints.

    Args:
        element: Paragraph XML element
        doc: Document object

    Returns:
        Extracted text with markdown-like formatting
    """
    from docx.oxml.ns import qn

    # Find the paragraph in the document
    for para in doc.paragraphs:
        if para._element == element:
            text = para.text.strip()

            # Check for heading styles
            style_name = para.style.name if para.style else ''
            if style_name.startswith('Heading'):
                # Convert to markdown-style heading
                level = style_name.replace('Heading ', '')
                try:
                    level_num = int(level)
                    prefix = '#' * min(level_num, 6)
                    return f"{prefix} {text}"
                except ValueError:
                    pass

            # Check for list items
            if para._element.xpath('.//w:numPr'):
                return f"• {text}"

            return text

    return ""


def _extract_table_text(element, doc) -> str:
    """
    Extract text from a table element in markdown format.

    Args:
        element: Table XML element
        doc: Document object

    Returns:
        Table text in markdown-like format
    """
    lines = []

    # Find the table in the document
    for table in doc.tables:
        if table._tbl == element:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    cells.append(cell_text)
                lines.append(' | '.join(cells))

            break

    if lines:
        return '\n'.join(lines)
    return ""
